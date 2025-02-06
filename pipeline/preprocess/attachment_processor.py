from pathlib import Path
from typing import Generator, Dict, Any
from dataclasses import dataclass
import re
from langchain_community.document_loaders import PyMuPDFLoader, TextLoader
from docx import Document
import openpyxl
import os

@dataclass
class ProcessedDocument:
    """Represents a processed document with its content and metadata"""
    content: str
    metadata: Dict[str, Any]

class DocumentProcessor:
    """Process different types of documents (PDF, DOCX, etc.) into text"""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'PDF Document',
        '.docx': 'Word Document',
        '.xlsx': 'Excel Spreadsheet',
        # Add more as we implement them
        # '.msg': 'Email Message'
    }

    # Headers that appear on every page that we want to remove
    REDUNDANT_PATTERNS = [
        r'Proprietary & Confidential\. Do Not Distribute\.',
        r'Page \d+ of \d+',
        r'^\s*\d+\s*$',  # Standalone page numbers
        r'(?i)overview$',  # Remove standalone "Overview" titles
        r'^[A-Z\s]+$',  # Standalone uppercase words like "HEALTHCARE"
    ]
    
    def __init__(self, base_dir: str = None):
        """Initialize the document processor
        
        Args:
            base_dir: Base directory for relative paths. If None, use absolute paths.
        """
        self.base_dir = Path(base_dir) if base_dir else None
    
    def _get_full_path(self, path: str) -> Path:
        """Get full path from potentially relative path"""
        path_obj = Path(path)
        if self.base_dir and not path_obj.is_absolute():
            return self.base_dir / path_obj
        return path_obj

    def _clean_text(self, text: str) -> str:
        """Clean and format extracted text."""
        # Standardize bullet points
        text = re.sub(r'[•●]', '•', text)
        
        # Remove redundant whitespace while preserving table structure
        lines = text.split('\n')
        cleaned_lines = []
        table_mode = False
        in_bullet_list = False
        
        for i, line in enumerate(lines):
            # Detect table structure based on multiple spaces or tabs
            if re.search(r'\t|\s{3,}', line):
                table_mode = True
            elif line.strip() == '':
                table_mode = False
            
            # Clean the line
            cleaned_line = line.strip()
            
            # Handle section headers
            if cleaned_line.isupper() and len(cleaned_line) > 3:
                if i > 0 and cleaned_lines:
                    cleaned_lines.append('')
                cleaned_lines.append(cleaned_line)
                if i < len(lines) - 1 and lines[i+1].strip():
                    cleaned_lines.append('')
                continue
            
            # Skip empty lines unless they're needed for structure
            if not cleaned_line:
                if table_mode or (i > 0 and i < len(lines) - 1 and 
                    not lines[i-1].strip().startswith('•') and
                    not lines[i+1].strip().startswith('•')):
                    cleaned_lines.append('')
                in_bullet_list = False
                continue
            
            # Preserve table formatting
            if table_mode:
                cleaned_line = line.rstrip()
            
            # Handle bullet points
            if cleaned_line.startswith('•'):
                if not in_bullet_list and cleaned_lines:
                    cleaned_lines.append('')
                cleaned_line = '• ' + cleaned_line[1:].strip()
                in_bullet_list = True
            else:
                in_bullet_list = False
            
            cleaned_lines.append(cleaned_line)
        
        # Join lines and clean up multiple line breaks
        text = '\n'.join(cleaned_lines)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Clean up table borders
        text = re.sub(r'[-]{3,}', '-' * 40, text)
        
        return text.strip()
    
    def _get_loader_from_extension(self, extension: str, path: str):
        """Get appropriate document loader based on file extension"""
        if extension == '.pdf':
            return PyMuPDFLoader(path)
        elif extension == '.docx':
            return self._load_docx(path)
        elif extension == '.xlsx':
            return self._load_xlsx(path)
        else:
            # Fallback to text loader
            return TextLoader(path)
    
    def _load_docx(self, path: str):
        """Load a DOCX file and return its content"""
        doc = Document(path)
        paragraphs = []
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append('\t'.join(row_text))
        
        return '\n'.join(paragraphs)

    def _load_xlsx(self, path: str):
        """Load an Excel file and return its content"""
        workbook = openpyxl.load_workbook(path, data_only=True)
        sheets_text = []
        
        for sheet in workbook.worksheets:
            # Get sheet name
            sheet_name = sheet.title
            sheets_text.append(f"Sheet: {sheet_name}")
            sheets_text.append("-" * 40)
            
            # Process cells
            rows_text = []
            max_col = sheet.max_column
            for row in sheet.iter_rows():
                row_text = []
                for cell in row[:max_col]:
                    value = cell.value
                    if value is not None:
                        # Convert numbers to strings with proper formatting
                        if isinstance(value, (int, float)):
                            if isinstance(value, int) or value.is_integer():
                                value = f"{int(value):,}"
                            else:
                                value = f"{value:,.2f}"
                        row_text.append(str(value))
                if any(text.strip() for text in row_text):
                    rows_text.append('\t'.join(row_text))
            
            if rows_text:
                sheets_text.extend(rows_text)
                sheets_text.append("")  # Add space between sheets
        
        return '\n'.join(sheets_text)

    def process_document(self, path: str) -> Generator[ProcessedDocument, None, None]:
        """Process a document and yield ProcessedDocument objects
        
        Args:
            path: Path to the document
            
        Yields:
            ProcessedDocument objects containing the document content and metadata
        
        Example:
            >>> processor = DocumentProcessor()
            >>> docs = list(processor.process_document("path/to/document.pdf"))
            >>> print(f"Content: {docs[0].content}")
            >>> print(f"Pages: {docs[0].metadata['total_pages']}")
        """
        full_path = self._get_full_path(path)
        
        if not full_path.exists():
            raise FileNotFoundError(f"File not found: {full_path}")
        
        extension = full_path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        
        loader = self._get_loader_from_extension(extension, str(full_path))
        
        if extension == '.docx':
            content = loader
        elif extension == '.xlsx':
            content = loader
        else:
            # Load all pages
            pages = list(loader.lazy_load())
            if not pages:
                return
            
            # Combine content from all pages
            content = '\n'.join(page.page_content for page in pages)
        
        content = self._clean_text(content)
        
        # Use metadata from first page, but include total pages
        metadata = {}
        if extension != '.docx' and extension != '.xlsx':
            metadata = pages[0].metadata.copy()
        metadata.update({
            'extension': extension,
            'source': str(full_path),
            'file_name': full_path.name,
            'document_type': self.SUPPORTED_EXTENSIONS[extension],
        })
        
        if extension != '.docx' and extension != '.xlsx':
            metadata['total_pages'] = len(pages)
        
        yield ProcessedDocument(
            content=content,
            metadata=metadata
        )

def test_pdf():
    """Test PDF document processing."""
    processor = DocumentProcessor()
    
    # Test with sample PDF files
    pdf_files = [
        'AlphaSense_Intro.pdf',
        'Brightwave_Report_on_NVIDIA_s_Blackwell_GPU.pdf',
        '242Q_ER_Deck_Eng_240725.pdf'
    ]
    
    for pdf_file in pdf_files:
        print(f"\nProcessing: {pdf_file}\n")
        print('-' * 100)
        for doc in processor.process_document(pdf_file):
            print(doc.content)
        print('-' * 100)

def test_docx():
    """Test DOCX document processing."""
    processor = DocumentProcessor()
    
    # Test with sample DOCX files
    docx_files = [
        "/home/minjae/email-rag/data/emails/Question_about_Your_AI_Needs_2025-01-26_06-51-54/message_1/LEI_Interview_Transcript_1_.docx",
        "/home/minjae/email-rag/data/emails/_Linqalpha_Constellation_Brands_1Q25_first_take_2025-01-26_06-51-54/message_1/SentinelOne_4Q24_earnings_call_transcript.docx"
    ]
    
    for docx_file in docx_files:
        print(f"\nProcessing: {docx_file}\n")
        print('-' * 100)
        for doc in processor.process_document(docx_file):
            print(doc.content)
        print('-' * 100)

def test_xlsx():
    """Test Excel document processing."""
    processor = DocumentProcessor()
    
    # Test with the problematic Excel file
    xlsx_file = "/home/minjae/email-rag/data/emails/CME_Group_December_2024_Volume_2025-01-26_06-51-14/message_1/cme_group_adv_rpc_oi_trend_Dec24_incl_cash_market.xlsx"
    
    print(f"\nProcessing: {xlsx_file}\n")
    print('-' * 100)
    try:
        docs = list(processor.process_document(xlsx_file))
        print(f"Successfully processed {len(docs)} documents")
        for doc in docs:
            print("\nDocument content:")
            print(doc.content[:500] + "..." if len(doc.content) > 500 else doc.content)
            print("\nMetadata:", doc.metadata)
    except Exception as e:
        print(f"Error processing file: {str(e)}")
    print('-' * 100)

if __name__ == '__main__':
    print("Testing Excel processing:")
    test_xlsx()
